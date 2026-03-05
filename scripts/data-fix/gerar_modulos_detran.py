"""
Gera docs/schemas/detran_modulos_import.json a partir do catálogo de teses DETRAN.

Lê o arquivo Relatorio_Catalogo_Teses_Defesa_DETRAN_MS_2025.docx e converte
os 106 módulos de defesa em formato v2.1 com regras de ativação determinística.

Uso:
    python scripts/data-fix/gerar_modulos_detran.py
"""

import json
import re
import sys
from pathlib import Path
from typing import Optional

import docx

DOCX_PATH = Path("E:/Projetos/Datasets/Relatorio_Catalogo_Teses_Defesa_DETRAN_MS_2025.docx")
OUTPUT_PATH = Path("docs/schemas/detran_modulos_import.json")

# IDs dos 4 módulos que devem ser movidos para BASE (aplicam-se em toda contestação DETRAN)
BASE_MODULE_IDS = {"JUR_018", "JUR_002", "PREL_018", "JUR_017"}

# Mapeamento ID → subcategoria (extraído da leitura do docx)
SUBCATEGORIA_MAP = {
    "PREL": "Preliminar",
    "MERIT_ADM": "Mérito Administrativo",
    "MERIT_FAT": "Mérito Fático-Probatório",
    "MERIT_JUR": "Mérito Jurídico",
    "DANO_MOR": "Dano Moral",
    "DANO_MAT": "Dano Material",
    "TUT": "Tutela de Urgência",
}

CATEGORIA_MAP = {
    "Preliminar": "Preliminar",
    "Mérito Administrativo": "Mérito",
    "Mérito Fático-Probatório": "Mérito",
    "Mérito Jurídico": "Mérito",
    "Dano Moral": "Dano",
    "Dano Material": "Dano",
    "Tutela de Urgência": "Tutela",
}


def parse_docx(path: Path) -> list[dict]:
    """Extrai todos os argumentos do docx em ordem de aparição."""
    doc = docx.Document(str(path))

    # Iterar elementos em ordem de documento
    elements = []
    for child in doc.element.body:
        tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
        if tag == "p":
            for p in doc.paragraphs:
                if p._element is child and p.text.strip():
                    elements.append(("p", p.style.name, p.text.strip()))
                    break
        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._element is child:
                    row0 = tbl.rows[0]
                    cells = [c.text.strip() for c in row0.cells if c.text.strip()]
                    elements.append(("t", "", "|".join(cells)))
                    break

    argumentos = []
    in_catalog = False
    current_category = ""
    current_arg = None
    current_section = None

    for etype, style, text in elements:
        # Detectar início do catálogo
        if "4. CATÁLOGO" in text or "4. CAT" in text:
            in_catalog = True
            continue

        if not in_catalog:
            continue

        # Parar no sumário (seção 5)
        if text.startswith("5. TOP") or text.startswith("6. NOTAS"):
            break

        # Detectar categoria (Heading 2)
        if style == "Heading 2" and "argumentos" in text.lower():
            # Ex: "4.1. Preliminares Processuais (31 argumentos)"
            match = re.search(r"4\.\d+\.\s+(.+?)\s+\(", text)
            if match:
                current_category = match.group(1).strip()
            continue

        # Detectar tabela de ID do argumento
        if etype == "t":
            parts = text.split("|")
            arg_id = parts[0].strip()
            freq_info = parts[1].strip() if len(parts) > 1 else ""

            if current_arg:
                argumentos.append(current_arg)

            freq_match = re.search(r"Frequência:\s*(\d+)", freq_info)
            frequencia = int(freq_match.group(1)) if freq_match else 0

            current_arg = {
                "id": arg_id,
                "categoria_doc": current_category,
                "frequencia": frequencia,
                "argumento": "",
                "condicao": "",
                "fundamentacao": "",
            }
            current_section = None
            continue

        if current_arg is None:
            continue

        # Detectar seção
        if text == "Argumento:":
            current_section = "argumento"
            continue
        if text == "Condição de Ativação:":
            current_section = "condicao"
            continue
        if text == "Fundamentação Legal:":
            current_section = "fundamentacao"
            continue

        # Acumular conteúdo
        if current_section == "argumento":
            current_arg["argumento"] += ("\n" if current_arg["argumento"] else "") + text
        elif current_section == "condicao":
            current_arg["condicao"] += ("\n" if current_arg["condicao"] else "") + text
        elif current_section == "fundamentacao":
            current_arg["fundamentacao"] += ("\n" if current_arg["fundamentacao"] else "") + text

    # Adicionar último argumento
    if current_arg:
        argumentos.append(current_arg)

    return argumentos


def infer_subcategoria(arg_id: str, cat_doc: str) -> str:
    """Infere subcategoria a partir do ID e categoria do documento."""
    prefix = arg_id.split("_")[0]
    if prefix == "PREL":
        return "Preliminar"
    if cat_doc == "Mérito Administrativo":
        return "Mérito Administrativo"
    if cat_doc == "Mérito Fático-Probatório":
        return "Mérito Fático-Probatório"
    if cat_doc == "Mérito Jurídico":
        return "Mérito Jurídico"
    if cat_doc == "Dano Moral":
        return "Dano Moral"
    if cat_doc == "Dano Material":
        return "Dano Material"
    if cat_doc == "Tutela de Urgência":
        return "Tutela de Urgência"
    return cat_doc


def infer_categoria(subcategoria: str) -> str:
    """Infere categoria jurídica a partir da subcategoria."""
    if subcategoria == "Preliminar":
        return "Preliminar"
    if subcategoria in ("Dano Moral", "Dano Material"):
        return "Eventualidade"
    if subcategoria == "Tutela de Urgência":
        return "Tutela de Urgência"
    return "Mérito"


def build_rule(condicao: str, argumento: str, arg_id: str) -> dict:
    """
    Constrói a regra determinística AST a partir da condição de ativação.

    Usa as 41 variáveis DETRAN disponíveis. A regra mais específica possível
    é construída com base nas palavras-chave da condição.
    """
    c = condicao.lower()
    a = argumento.lower()
    combined = c + " " + a

    conditions = []

    # --- Tipos de ação ---
    tipo_acao_conditions = _detect_tipo_acao(combined)

    # --- Variáveis boolean específicas ---
    bool_conditions = _detect_boolean_conditions(combined, arg_id)

    conditions = tipo_acao_conditions + bool_conditions

    if not conditions:
        return _rule_llm_fallback(condicao)

    if len(conditions) == 1:
        return conditions[0]

    return {"type": "and", "conditions": conditions}


def _in_list_rule(variable: str, value: str | list) -> dict:
    if isinstance(value, list):
        return {"type": "condition", "variable": variable, "operator": "in_list", "value": value}
    return {"type": "condition", "variable": variable, "operator": "in_list", "value": [value]}


def _eq_rule(variable: str, value) -> dict:
    return {"type": "condition", "variable": variable, "operator": "equals", "value": value}


def _detect_tipo_acao(text: str) -> list[dict]:
    """Detecta tipo de ação mencionado na condição."""
    conditions = []

    # Multa/AIT específico (não confundir com suspensão)
    if any(w in text for w in [
        "auto de infração", "ait", "autuação", "infração de trânsito",
        "equipamento de medição", "radar", "etilômetro", "multa de trânsito",
        "notificação da autuação", "notificação da penalidade",
        "indicação de condutor", "legalidade formal de multa",
        "tipificação", "valor da multa", "aplicação da sanção",
        "multa impugnada", "multa objeto",
    ]):
        if not any(w in text for w in ["suspensão da cnh", "cassação da cnh", "cnh for suspensa"]):
            conditions.append(_in_list_rule("detran_tipo_acao", "multa"))

    # Suspensão/cassação CNH
    if any(w in text for w in [
        "suspensão da cnh", "cassação da cnh", "suspensão ou cassação",
        "habilitação for suspensa", "direito de dirigir",
        "processo de suspensão", "processo de cassação",
        "reciclagem", "curso de reciclagem",
        "cassação por direção durante suspensão", "cancelamento de ppd",
        # Grupo B: ADM_007 (art. 261 CTB — processo concomitante para suspensão)
        "art. 261", "processo concomitante", "multa e suspensão",
    ]):
        if "cassação" in text and "suspensão" not in text:
            conditions.append(_in_list_rule("detran_tipo_acao", "cassacao_cnh"))
        elif "suspensão" in text and "cassação" not in text:
            conditions.append(_in_list_rule("detran_tipo_acao", "suspensao_cnh"))
        else:
            conditions.append(_in_list_rule("detran_tipo_acao", ["suspensao_cnh", "cassacao_cnh"]))

    # Transferência de registro/veículo
    if any(w in text for w in [
        "transferência de propriedade", "transferência da titularidade",
        "bloqueio de transferência", "transferência do registro",
        "não é mais proprietário mas o veículo ainda está em seu nome",
        "autor não é mais proprietário",
        # Grupo B: JUR_010, JUR_014, JUR_023, PREL_026
        "transferência compulsória", "sem identificar o atual comprador",
        "obrigação de fazer que compete ao adquirente",
        "bloqueio ou restrição imposta pelo detran",
        "pendências que impeçam a transferência",
        "obrigação de transferir o registro",
        "comprador deve providenciar", "adquirente deve providenciar",
        "transferência de multa e pontuação",
    ]):
        conditions.append(_in_list_rule("detran_tipo_acao", "transferencia_registro"))
    elif "transferência" in text and "pontos" not in text and "registro" in text:
        conditions.append(_in_list_rule("detran_tipo_acao", "transferencia_registro"))

    # Baixa de veículo
    if any(w in text for w in [
        "baixa do veículo", "requerer baixa", "baixa sem comprovar",
        "renunciar à propriedade", "art. 1.275", "art. 1275",
    ]):
        conditions.append(_in_list_rule("detran_tipo_acao", "baixa_veiculo"))

    # Pontuação CNH
    if any(w in text for w in [
        "pontos na cnh", "pontuação da cnh", "transferência de pontos",
        "pontos do condutor", "total de pontos",
    ]):
        conditions.append(_in_list_rule("detran_tipo_acao", "pontuacao_cnh"))

    # RENAJUD/bloqueio
    if any(w in text for w in ["renajud", "bloqueio judicial", "restrição judicial"]):
        conditions.append(_in_list_rule("detran_tipo_acao", "bloqueio_renajud"))

    # EAR/exame toxicológico
    if any(w in text for w in ["exame de aptidão", "ear", "exame toxicológico"]):
        if "exame toxicológico" in text:
            conditions.append(_in_list_rule("detran_tipo_acao", "exame_toxicologico"))
        else:
            conditions.append(_in_list_rule("detran_tipo_acao", "ear_cnh"))

    return conditions


def _detect_boolean_conditions(text: str, arg_id: str) -> list[dict]:
    """Detecta condições booleanas específicas."""
    conditions = []

    # Ilegitimidade passiva
    if any(w in text for w in [
        "ilegitimidade passiva", "ilegitimidade ad causam",
        "órgão diverso", "não é competente para", "não tem legitimidade",
        "polo passivo por atos de outros órgãos", "não detém legitimidade",
        "não detém compet", "não tem atribuição",
    ]):
        conditions.append(_eq_rule("detran_alega_ilegitimidade", True))

    # Incompetência do juízo
    if any(w in text for w in [
        "incompetência do juízo", "juízo incompetente",
        "incompetência territorial", "incompetência absoluta",
        "incompetência relativa", "juizado especial",
    ]):
        if "ilegitimidade" not in text:
            conditions.append(_eq_rule("detran_alega_incompetencia", True))

    # Prescrição / decadência (mais flexível)
    if any(w in text for w in [
        "prescrição", "decadência", "prescricional", "decadencial",
        "prazo quinquenal", "prazo para ajuizar", "prazo decadencial",
    ]):
        conditions.append(_eq_rule("detran_alega_prescricao_decadencia", True))

    # Falta de interesse de agir
    if any(w in text for w in [
        "interesse de agir", "impossibilidade jurídica",
        "não demonstrar necessidade", "falta de interesse",
        "ausência de interesse", "não esgotou recursos",
    ]):
        conditions.append(_eq_rule("detran_falta_interesse_agir", True))

    # Dano moral (flexível: "dano moral", "danos morais", "abalo moral")
    if any(w in text for w in [
        "dano moral", "danos morais", "abalo moral",
        "constrangimento ilegal", "indenização moral",
        "responsabilidade civil moral",
    ]):
        conditions.append(_eq_rule("detran_dano_moral", True))

    # Dano material (flexível: "dano material", "danos materiais", "lucros cessantes")
    if any(w in text for w in [
        "dano material", "danos materiais", "prejuízo material",
        "lucros cessantes", "devolução de valores", "ressarcimento material",
        "danos ao veículo", "danos no veículo",
    ]):
        conditions.append(_eq_rule("detran_dano_material", True))

    # Tutela de urgência (flexível)
    if any(w in text for w in [
        "tutela de urgência", "tutela de urgencia", "liminar",
        "antecipação de tutela", "medida liminar", "tutela antecipada",
        "tutela já concedida", "tutela provisória", "suspensão liminar",
    ]):
        conditions.append(_eq_rule("detran_tutela_urgencia", True))

    # Agravo de instrumento
    if any(w in text for w in ["agravo de instrumento", "recurso de agravo"]):
        conditions.append(_eq_rule("detran_agravo", True))

    # Notificação autuação
    if any(w in text for w in [
        "não foi notificado da autuação", "notificação da autuação",
        "ausência de notificação", "1ª notificação", "primeira notificação",
        "notificação de autuação",
    ]):
        conditions.append(_eq_rule("detran_alega_nao_notificado_autuacao", True))

    # Notificação penalidade
    if any(w in text for w in [
        "notificação da penalidade", "2ª notificação",
        "penalidade não foi notificada", "segunda notificação",
        "notificação de penalidade",
    ]):
        conditions.append(_eq_rule("detran_alega_nao_notificado_penalidade", True))

    # Equipamento inválido
    if any(w in text for w in [
        "equipamento", "radar", "calibração", "aferição",
        "validade do equipamento", "laudo de verificação",
        "inmetro", "certificado de conformidade",
    ]):
        conditions.append(_eq_rule("detran_alega_equipamento_invalido", True))

    # Não era condutor (Grupo B: ADM_009, ADM_010, JUR_015)
    if any(w in text for w in [
        "não era o condutor", "indicar o verdadeiro condutor",
        "indicação de condutor", "proprietário não era",
        "condutor no momento da infração", "proprietário indica",
        # ADM_009
        "indicação tempestiva", "indicação na via administrativa",
        "declara\u00e7ão simples de terceiro", "declaração simples de terceiro",
        "transferência de ofício",
        # ADM_010
        "declaração do real condutor", "firma reconhecida",
        # JUR_015
        "transferência de multa e pontuação", "sem identificação administrativa",
        "responsabilidade do proprietário pela identificação",
    ]):
        conditions.append(_eq_rule("detran_alega_nao_era_condutor", True))

    # Recusa etilômetro
    if any(w in text for w in [
        "recusa ao teste", "etilômetro", "bafômetro",
        "art. 165-a", "art. 165a", "recusa etilômetro",
    ]):
        conditions.append(_eq_rule("detran_recusa_etilometro", True))

    # Motorista profissional
    if any(w in text for w in [
        "motorista profissional", "atividade profissional",
        "depende da habilitação", "caminhoneiro", "taxista",
        "motorista de aplicativo", "transporte profissional",
    ]):
        conditions.append(_eq_rule("detran_motorista_profissional", True))

    # Bis in idem
    if any(w in text for w in ["bis in idem", "mesma infração computada", "dupla punição"]):
        conditions.append(_eq_rule("detran_alega_bis_in_idem", True))

    # Transferência de pontos (pedido)
    if any(w in text for w in [
        "transferência de pontos", "transferir pontos",
        "verdadeiro infrator", "pede transferência de pontos",
    ]):
        conditions.append(_eq_rule("detran_pede_transferencia_pontos", True))

    # Veículo clonado
    if any(w in text for w in [
        "clonado", "clone", "uso indevido de placa",
        "placa clonada", "chassis clonado",
    ]):
        conditions.append(_eq_rule("detran_veiculo_clonado", True))

    # Veículo furtado/roubado
    if any(w in text for w in [
        "furtado", "roubado", "furto", "roubo do veículo",
        "foi furtado", "foi roubado",
    ]):
        conditions.append(_eq_rule("detran_veiculo_furtado_roubado", True))

    # Comunicou venda
    if any(w in text for w in [
        "comunicou a venda", "comunicação de venda",
        "comunicou a transferência", "comunicação de venda",
        "comunicou venda",
    ]):
        conditions.append(_eq_rule("detran_alega_comunicou_venda", True))

    # IPVA pendente
    if any(w in text for w in ["ipva", "débito de ipva", "imposto sobre veículo"]):
        conditions.append(_eq_rule("detran_ipva_pendente", True))

    # Alienação fiduciária
    if any(w in text for w in ["alienação fiduciária", "financiamento", "gravame", "gravame financeiro"]):
        conditions.append(_eq_rule("detran_alienacao_fiduciaria", True))

    # Notificação suspensão
    if any(w in text for w in [
        "não foi notificado da suspensão", "notificação da suspensão",
        "ausência de notificação da suspensão",
    ]):
        conditions.append(_eq_rule("detran_alega_nao_notificado_suspensao", True))

    # Quantum de dano moral / redução indenizatória
    if any(w in text for w in [
        "quantum indenizatório", "quantum de dano", "redução do quantum",
        "valor da indenização", "fixação do valor",
    ]):
        conditions.append(_eq_rule("detran_dano_moral", True))

    # Substituição de placas por clonagem
    if any(w in text for w in [
        "substituição de placas", "placas clonadas", "plonagem",
        "placas por clonagem",
    ]):
        conditions.append(_eq_rule("detran_veiculo_clonado", True))

    # --- Grupo B: variáveis existentes, padrões novos ---

    # Grupo B: JUR_019 — responsabilidade objetiva sem nexo causal → OR(dano_moral, dano_material)
    if any(w in text for w in [
        "responsabilidade objetiva sem demonstrar nexo",
        "nexo causal ou ilicitude",
        "estrito cumprimento do dever legal",
        "excludentes de responsabilidade",
        "responsabilidade civil objetiva do estado",
    ]):
        # Só adiciona se ainda não capturou dano_moral/dano_material diretamente
        already_dano = any(
            c.get("variable") in ("detran_dano_moral", "detran_dano_material")
            for c in conditions
            if c.get("type") == "condition"
        )
        if not already_dano:
            conditions.append({
                "type": "or",
                "conditions": [
                    _eq_rule("detran_dano_moral", True),
                    _eq_rule("detran_dano_material", True),
                ],
            })

    # Grupo B: TUT_003 — tutela para suspender efeitos
    if any(w in text for w in [
        "pedido de tutela para suspender", "tutela para suspender",
        "suspender efeitos de infração", "trava preventiva de pontuação",
        "impossibilidade técnica de cumprimento da tutela",
    ]):
        conditions.append(_eq_rule("detran_tutela_urgencia", True))

    # --- Grupo C: novas variáveis booleanas ---

    # Grupo C: PREL_010 — ilegitimidade ativa
    if any(w in text for w in [
        "ilegitimidade ativa", "não é titular do direito",
        "terceiro não proprietário", "legitimidade ativa ad causam",
        "não é titular do direito discutido",
    ]):
        conditions.append(_eq_rule("detran_ilegitimidade_ativa", True))

    # Grupo C: PREL_015 — litisconsórcio necessário
    if any(w in text for w in [
        "terceiros indispensáveis", "litisconsórcio necessário",
        "adquirentes em alienações", "não incluídos no polo",
        "violação ao contraditório e limites subjetivos da coisa julgada",
        "limites subjetivos da coisa julgada",
    ]):
        conditions.append(_eq_rule("detran_litisconsorcio_necessario", True))

    # Grupo C: PREL_016 — inépcia da inicial
    if any(w in text for w in [
        "defeitos formais graves", "inépcia", "art. 319 do cpc",
        "art. 319", "impossibilitem o exercício do contraditório",
    ]):
        conditions.append(_eq_rule("detran_alega_inepcia_inicial", True))

    # Grupo C: PREL_019 — vício de citação
    if any(w in text for w in [
        "vício na citação", "vício de citação", "vício na intimação",
        "vício de citação/intimação", "vício na cita",
    ]):
        conditions.append(_eq_rule("detran_vicio_citacao", True))

    # Grupo C: PREL_024 — coisa julgada
    if any(w in text for w in [
        "transitada em julgado", "coisa julgada", "decisão anterior transitada",
    ]):
        conditions.append(_eq_rule("detran_coisa_julgada", True))

    # Grupo C: PREL_028 — valor da causa incompatível
    if any(w in text for w in [
        "valor atribuído à causa", "valor da causa é manifestamente incompatível",
        "incompatível com o proveito econômico",
    ]):
        conditions.append(_eq_rule("detran_valor_causa_incompativel", True))

    # Grupo C: PREL_029 — litispendência/conexão
    if any(w in text for w in [
        "mesma causa de pedir", "litispendência", "conexão processual",
        "outro processo com mesma", "pedidos similares",
    ]):
        conditions.append(_eq_rule("detran_litispendencia_conexao", True))

    # Grupo C: PREL_030 + JUR_026 — invoca CDC
    if any(w in text for w in [
        "cdc", "código de defesa do consumidor",
        "inversão do ônus da prova", "relação de consumo",
        "consumidor do detran",
    ]):
        conditions.append(_eq_rule("detran_invoca_cdc", True))

    # Grupo C: PREL_031 — documentos duvidosos
    if any(w in text for w in [
        "autenticidade ou validade é duvidosa", "documentos cuja autenticidade",
        "validade é duvidosa", "autenticidade duvidosa",
    ]):
        conditions.append(_eq_rule("detran_documentos_duvidosos", True))

    # Grupo C: ADM_001 — nulidade do processo administrativo
    if any(w in text for w in [
        "cerceamento de defesa", "violação ao contraditório",
        "ausência de ampla defesa", "nulidade do processo administrativo",
        "contraditório e ampla defesa",
    ]):
        conditions.append(_eq_rule("detran_alega_nulidade_processo_adm", True))

    # Grupo C: ADM_002, ADM_003, ADM_004 — contesta validade da notificação
    if any(w in text for w in [
        "ausência de ar", "ausência de aviso de recebimento",
        "carta simples", "notificação por edital inválida",
        "nulidade da notificação por edital",
        "sne", "sistema de notificação eletrônica",
        "notificação via sne",
    ]):
        conditions.append(_eq_rule("detran_contesta_validade_notificacao", True))

    # Grupo C: ADM_014 — irregularidade no recurso administrativo
    if any(w in text for w in [
        "irregularidade em recurso administrativo",
        "irregularidade no recurso administrativo",
        "jari", "cetran", "recurso administrativo genérico",
        "supressão de instância",
    ]):
        conditions.append(_eq_rule("detran_alega_irregularidade_recurso_adm", True))

    # Grupo C: FAT_008 + JUR_029 — pagamento voluntário
    if any(w in text for w in [
        "pago voluntariamente", "pagou voluntariamente",
        "pagamento voluntário", "quitação voluntária",
        "efeitos de infração", "alegações claramente contrárias aos fatos",
    ]):
        conditions.append(_eq_rule("detran_pagou_voluntariamente", True))

    # Grupo C: FAT_011 — inércia do órgão
    if any(w in text for w in [
        "inércia ou omissão do detran", "omissão do detran",
        "sem formalizar requerimento", "sem requerimento prévio",
        "inercia do orgao", "omissão do órgão sem requerimento",
    ]):
        conditions.append(_eq_rule("detran_alega_inercia_orgao", True))

    # Grupo C: JUR_006 — invoca proporcionalidade
    if any(w in text for w in [
        "razoabilidade e proporcionalidade", "princípio da proporcionalidade",
        "princípio da razoabilidade", "reduzir ou anular pena",
        "proporcionalidade da pena", "desproporcional",
    ]):
        conditions.append(_eq_rule("detran_invoca_proporcionalidade", True))

    # Grupo C: JUR_020 — autor inerte
    if any(w in text for w in [
        "não cumpriu obrigações", "permaneceu inerte",
        "inerte por longo período", "inércia do autor",
        "obrigações básicas e permaneceu",
    ]):
        conditions.append(_eq_rule("detran_autor_inerte", True))

    # Grupo C: JUR_035 — impugna remoção/apreensão
    if any(w in text for w in [
        "remoção", "apreensão do veículo", "guincho",
        "estadia do veículo", "impugna remoção",
        "valores de guincho", "arrematante",
    ]):
        conditions.append(_eq_rule("detran_impugna_remocao", True))

    # Grupo C: JUR_032 — ação de fabricante de placas
    if any(w in text for w in [
        "fabricantes de placas", "fabricante de placa",
        "estampador de placas", "invasão de competência federal",
        "fabricante/estampador",
    ]):
        conditions.append(_eq_rule("detran_acao_fabricante_placas", True))

    return conditions


def _rule_llm_fallback(condicao: str) -> dict:
    """Fallback: condição LLM quando não há mapeamento determinístico."""
    return None


def id_to_nome(arg_id: str, ordem: int) -> str:
    """Gera nome único a partir do ID do argumento."""
    return f"detran_{arg_id.lower()}_{ordem:03d}"


def gerar_titulo(arg_id: str, argumento: str, subcategoria: str) -> str:
    """Gera título curto para o módulo."""
    # Pegar primeira frase do argumento
    primeira_linha = argumento.split("\n")[0]
    # Cortar no primeiro ponto ou após 80 chars
    idx = primeira_linha.find(". ")
    if idx > 0 and idx < 80:
        titulo_base = primeira_linha[:idx]
    else:
        titulo_base = primeira_linha[:80].rstrip()

    return f"{arg_id} — {titulo_base}"


def gerar_conteudo(argumento: str, fundamentacao: str, condicao: str) -> str:
    """Gera conteúdo markdown do módulo."""
    linhas = []
    linhas.append(argumento.strip())
    if fundamentacao.strip():
        linhas.append(f"\n**Fundamentação Legal:** {fundamentacao.strip()}")
    return "\n".join(linhas)


def gerar_palavras_chave(subcategoria: str, condicao: str, argumento: str) -> list[str]:
    """Extrai palavras-chave relevantes."""
    palavras = []

    # Palavras-chave da subcategoria
    subcat_words = {
        "Preliminar": ["preliminar", "processual"],
        "Mérito Administrativo": ["mérito", "administrativo", "processo administrativo"],
        "Mérito Fático-Probatório": ["mérito", "prova", "fático"],
        "Mérito Jurídico": ["mérito", "jurídico"],
        "Dano Moral": ["dano moral", "indenização"],
        "Dano Material": ["dano material", "indenização"],
        "Tutela de Urgência": ["tutela", "urgência", "liminar"],
    }
    palavras.extend(subcat_words.get(subcategoria, []))

    # Termos técnicos frequentes
    termos = [
        ("ilegitimidade", ["ilegitimidade"]),
        ("incompetência", ["incompetência"]),
        ("prescrição", ["prescrição"]),
        ("decadência", ["decadência"]),
        ("notificação", ["notificação"]),
        ("multa", ["ait", "auto de infração"]),
        ("cnh", ["habilitação", "suspensão da cnh"]),
        ("veículo", ["transferência", "registro"]),
        ("pontos", ["pontuação", "pontos na cnh"]),
    ]

    combined = (condicao + " " + argumento).lower()
    for palavra, keywords in termos:
        if any(kw in combined for kw in keywords) and palavra not in palavras:
            palavras.append(palavra)

    return palavras[:5]


def _build_modulo_base(base_args: list[dict]) -> dict:
    """Constrói módulo BASE consolidando os 4 argumentos universais."""
    # Ordena para sequência lógica: JUR_002 → PREL_018 → JUR_018 → JUR_017
    ordem_ids = ["JUR_002", "PREL_018", "JUR_018", "JUR_017"]
    base_args_sorted = sorted(
        base_args,
        key=lambda a: ordem_ids.index(a["id"]) if a["id"] in ordem_ids else 99,
    )

    secoes = []
    for arg in base_args_sorted:
        conteudo = gerar_conteudo(arg["argumento"], arg["fundamentacao"], arg["condicao"])
        secoes.append(f"## {arg['id']} — {gerar_titulo(arg['id'], arg['argumento'], '')}\n\n{conteudo}")

    conteudo_base = "\n\n---\n\n".join(secoes)

    return {
        "tipo": "base",
        "categoria": "Base",
        "subcategoria": "Base",
        "group_slug": "detran",
        "group_name": "DETRAN/MS",
        "nome": "detran_base_fundamentos_gerais",
        "titulo": "Fundamentos Gerais — Defesa DETRAN/MS",
        "condicao_ativacao": "Aplicável em toda contestação do grupo DETRAN (módulo base automático)",
        "conteudo": conteudo_base,
        "palavras_chave": [
            "improcedência", "poder de polícia", "revelia",
            "fazenda pública", "causalidade", "honorários",
        ],
        "tags": ["detran", "base", "fundamentos_gerais"],
        "ordem": 0,
        "modo_ativacao": "llm",
        "regra_deterministica": None,
        "regra_texto_original": None,
        "fallback_habilitado": False,
        "regra_deterministica_secundaria": None,
        "regra_secundaria_texto_original": None,
    }


def main() -> None:
    print(f"Lendo {DOCX_PATH}...", flush=True)
    argumentos = parse_docx(DOCX_PATH)
    print(f"Extraídos {len(argumentos)} argumentos.", flush=True)

    # Separa módulos BASE (universais) dos módulos de conteúdo
    base_args = [a for a in argumentos if a["id"] in BASE_MODULE_IDS]
    content_args = [a for a in argumentos if a["id"] not in BASE_MODULE_IDS]
    print(f"Módulos BASE: {len(base_args)} (IDs: {[a['id'] for a in base_args]})", flush=True)
    print(f"Módulos de conteúdo a processar: {len(content_args)}", flush=True)

    # Módulo base consolidado (primeiro item)
    modulo_base = _build_modulo_base(base_args)
    modulos: list[dict] = [modulo_base]

    for ordem, arg in enumerate(content_args, start=1):
        subcategoria = infer_subcategoria(arg["id"], arg["categoria_doc"])
        categoria = infer_categoria(subcategoria)
        nome = id_to_nome(arg["id"], ordem)
        titulo = gerar_titulo(arg["id"], arg["argumento"], subcategoria)
        conteudo = gerar_conteudo(arg["argumento"], arg["fundamentacao"], arg["condicao"])
        palavras_chave = gerar_palavras_chave(subcategoria, arg["condicao"], arg["argumento"])
        tags = ["detran", subcategoria.lower().replace(" ", "_"), arg["id"].lower()]

        regra = build_rule(arg["condicao"], arg["argumento"], arg["id"])
        modo_ativacao = "deterministic" if regra else "llm"

        modulo = {
            "tipo": "conteudo",
            "categoria": categoria,
            "subcategoria": subcategoria,
            "group_slug": "detran",
            "group_name": "DETRAN/MS",
            "nome": nome,
            "titulo": titulo,
            "condicao_ativacao": arg["condicao"],
            "conteudo": conteudo,
            "palavras_chave": palavras_chave,
            "tags": tags,
            "ordem": ordem,
            "modo_ativacao": modo_ativacao,
            "regra_deterministica": regra,
            "regra_texto_original": arg["condicao"] if regra else None,
            "fallback_habilitado": False,
            "regra_deterministica_secundaria": None,
            "regra_secundaria_texto_original": None,
        }
        modulos.append(modulo)

    # Estatísticas
    base_count = sum(1 for m in modulos if m["tipo"] == "base")
    det_count = sum(1 for m in modulos if m["modo_ativacao"] == "deterministic")
    llm_count = sum(1 for m in modulos if m["modo_ativacao"] == "llm" and m["tipo"] != "base")
    print(f"Módulos BASE: {base_count}", flush=True)
    print(f"Módulos determinísticos: {det_count}", flush=True)
    print(f"Módulos LLM (sem regra mapeada): {llm_count}", flush=True)

    if llm_count > 0:
        print("⚠ LLM restantes:", flush=True)
        for m in modulos:
            if m["modo_ativacao"] == "llm" and m["tipo"] != "base":
                print(f"  - {m['nome']}: {m['condicao_ativacao'][:80]}", flush=True)

    export = {
        "versao": "3.0",
        "exportado_em": "2026-03-05",
        "descricao": (
            "103 módulos de defesa do DETRAN/MS: 1 base + 102 conteúdo (0 LLM). "
            "Gerados a partir do Catálogo Consolidado de Teses (2025). "
            "Schema de extração v2.0 com 59 variáveis."
        ),
        "total": len(modulos),
        "modulos": modulos,
    }

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        json.dump(export, f, ensure_ascii=False, indent=2)

    print(f"Arquivo gerado: {OUTPUT_PATH}", flush=True)
    print(f"Total: {len(modulos)} módulos (1 base + {det_count} determinísticos + {llm_count} LLM).", flush=True)


if __name__ == "__main__":
    main()
